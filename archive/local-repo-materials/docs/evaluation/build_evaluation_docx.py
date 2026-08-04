from __future__ import annotations

import csv
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "evaluation"
DOCX_PATH = OUT_DIR / "OAlpha_Evaluation_Summary_and_User_Testing_Readout_2026-07-08.docx"
CHART_PATH = OUT_DIR / "OAlpha_Evaluation_Data_Snapshot_2026-07-08.png"
CSV_PATH = OUT_DIR / "solution_suitability_metrics_2026-07-05.csv"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(100, 116, 139)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF7ED"
PALE_GREEN = "ECFDF5"
PALE_RED = "FEF2F2"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, col_widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(col_widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(w))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(col_widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(col_widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def style_cell_text(cell, size=9, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        for run in p.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)


def add_body_para(doc, text, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_callout(doc, title, body, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=INK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_simple_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], header_fill)
        style_cell_text(hdr[i], size=font_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            style_cell_text(cells[i], size=font_size)
    set_table_width(table, widths)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        keep_row_together(row)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if p.runs:
            p.runs[0].text = item
            set_run_font(p.runs[0], size=11, color=INK)
        else:
            run = p.add_run(item)
            set_run_font(run, size=11, color=INK)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def set_headers(doc):
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = ""
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("O(Alpha) Evaluation Readout")
    set_run_font(run, size=9, color=MUTED)
    footer = section.footer
    footer.paragraphs[0].text = ""
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Private product planning artifact - simulated proxy data clearly labelled")
    set_run_font(run, size=8.5, color=MUTED)


def load_metrics():
    with CSV_PATH.open(newline="") as f:
        return list(csv.DictReader(f))


def draw_chart():
    img = Image.new("RGB", (1800, 1050), "#f8fafc")
    draw = ImageDraw.Draw(img)

    def font(size, bold=False):
        candidates = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
            "/System/Library/Fonts/Helvetica.ttc",
        ]
        for path in candidates:
            try:
                return ImageFont.truetype(path, size)
            except Exception:
                continue
        return ImageFont.load_default()

    f_title = font(42, True)
    f_h = font(26, True)
    f_label = font(20, True)
    f_body = font(18, False)
    f_small = font(15, False)
    f_value = font(32, True)

    def rounded(xy, fill, outline=None, width=2, radius=22):
        draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)

    draw.text((70, 60), "Evaluation Data Snapshot", fill="#0f172a", font=f_title)
    draw.text((70, 112), "Simulated proxy evaluation, N=10. Not real survey or actual-user statistics.", fill="#475569", font=f_body)
    rounded((1320, 56, 1718, 112), "#fff7ed", "#fed7aa", radius=28)
    draw.text((1350, 74), "SIMULATED PROXY DATA", fill="#9a3412", font=f_label)

    rounded((70, 170, 570, 470), "#ffffff", "#dbe4ef")
    draw.text((105, 205), "Method Status", fill="#0f172a", font=f_h)
    status = [("Completed defensible", 3, 7, "#2563eb"), ("Ready to field later", 4, 7, "#64748b")]
    for i, (label, value, maxv, color) in enumerate(status):
        y = 265 + i * 95
        draw.text((105, y), label, fill="#334155", font=f_label)
        rounded((105, y + 35, 470, y + 68), "#e2e8f0", None, radius=16)
        rounded((105, y + 35, 105 + int(365 * value / maxv), y + 68), color, None, radius=16)
        draw.text((485, y + 34), f"{value}/{maxv}", fill="#0f172a", font=f_label)

    rounded((70, 510, 570, 920), "#ffffff", "#dbe4ef")
    draw.text((105, 545), "Average Ratings", fill="#0f172a", font=f_h)
    ratings = [("Ease", 3.4, "#16a34a"), ("Confidence", 2.9, "#d97706"), ("Transparency", 2.5, "#e11d48")]
    for i, (label, value, color) in enumerate(ratings):
        y = 620 + i * 85
        draw.text((105, y), label, fill="#334155", font=f_label)
        rounded((255, y, 485, y + 34), "#e2e8f0", None, radius=17)
        rounded((255, y, 255 + int(230 * value / 5), y + 34), color, None, radius=17)
        draw.text((505, y - 1), f"{value:.1f}/5", fill="#0f172a", font=f_label)

    rounded((620, 170, 1730, 920), "#ffffff", "#dbe4ef")
    draw.text((655, 205), "Repeated Themes", fill="#0f172a", font=f_h)
    themes = [
        ("Settings lock felt reassuring", 10, "#0f766e"),
        ("Clearer paper-only messaging needed", 10, "#2563eb"),
        ("Validation evidence should be visible in UI", 9, "#7c3aed"),
        ("Quant terms or strategy names unclear", 8, "#d97706"),
        ("Freshness or health indicators needed", 7, "#e11d48"),
        ("Trade rationale drilldowns needed", 6, "#64748b"),
    ]
    for i, (label, value, color) in enumerate(themes):
        y = 275 + i * 92
        draw.text((655, y), label, fill="#334155", font=f_label)
        rounded((1110, y - 3, 1600, y + 34), "#e2e8f0", None, radius=18)
        rounded((1110, y - 3, 1110 + int(490 * value / 10), y + 34), color, None, radius=18)
        draw.text((1620, y - 2), f"{value}/10", fill="#0f172a", font=f_label)

    draw.line((70, 965, 1730, 965), fill="#cbd5e1", width=2)
    draw.text((70, 990), "Interpretation: users can complete the core flow, but trust depends on clearer paper-only limits, visible strategy evidence, and operational freshness indicators.", fill="#475569", font=f_small)
    img.save(CHART_PATH)


def build_doc():
    draw_chart()
    doc = Document()
    setup_styles(doc)
    set_headers(doc)

    # Masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("O(Alpha)")
    set_run_font(run, size=13, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Evaluation Summary and User Testing Readout")
    set_run_font(run, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Solution suitability, simulated proxy user feedback, and next research plan")
    set_run_font(run, size=13, color=MUTED)

    meta = [
        ("Prepared for", "O(Alpha) product and submission planning"),
        ("Date", "8 July 2026"),
        ("Evidence status", "Expert/self evaluation, heuristic walkthrough, and simulated proxy persona review completed"),
        ("Integrity note", "No public launch, real user survey, or actual participant interview has been conducted yet"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=10.5, color=INK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=INK)

    add_callout(
        doc,
        "Recommended claim for the submission",
        "The team used expert/self evaluation, a cognitive walkthrough / heuristic evaluation, and a simulated 10-person proxy persona review. Real participant interviews, low-fidelity tests, surveys, and high-fidelity usability tests are prepared as privacy-safe protocols but should not be marked as completed until fielded with real users.",
        fill=PALE_YELLOW,
    )

    doc.add_heading("1. Executive Summary", level=1)
    add_body_para(
        doc,
        "O(Alpha)'s core paper-trading workflow is understandable and safety-aware: users select a risk profile, run and accept a backtest, launch or terminate a paper agent, and are prevented from changing settings while a portfolio agent is active. The largest product gap is trust calibration. Users need stronger paper-only messaging, visible strategy evidence, and operational freshness indicators before they can confidently judge the system.",
    )
    add_bullets(
        doc,
        [
            "Core workflow health: usable, with average simulated ease of 3.4/5.",
            "Trust confidence: moderate at 2.9/5, mainly because evidence is not visible enough in the UI.",
            "Transparency: weakest area at 2.5/5; validation proof is stronger in repo artifacts than in the app screens.",
            "Safety: active-agent settings lock was the most consistently positive signal, raised by 10/10 proxy personas.",
            "Top improvement: repeat paper-only / no brokerage orders messaging near Launch Agent, P&L, backtest acceptance, and settings.",
        ]
    )

    doc.add_page_break()
    doc.add_heading("2. Evaluation Methods", level=1)
    method_rows = [
        ("Expert / self evaluation", "Completed", "Yes", "Review of docs, frontend flows, safety behavior, and validation story."),
        ("Cognitive walkthrough / heuristic evaluation", "Completed", "Yes", "Task walkthrough across onboarding, dashboard, launch/terminate, settings, and evidence discovery."),
        ("Simulated user focus group", "Completed, simulated", "Yes, labelled simulated", "Ten proxy personas; six from parallel subagent reviews and four local synthetic personas."),
        ("Actual user focus group / interview", "Not conducted", "No", "Private 10-person protocol prepared for future fielding."),
        ("Low-fidelity usability testing", "Not conducted", "No", "Figma/wireframe task protocol prepared."),
        ("Survey of potential users", "Not fielded", "No", "50-person survey plan and item bank prepared."),
        ("High-fidelity usability testing", "Not conducted with real users", "No", "Seeded-account working prototype protocol prepared."),
    ]
    add_simple_table(
        doc,
        ["Method", "Status", "Select now?", "Evidence or next step"],
        method_rows,
        [2500, 1600, 1400, 3860],
        font_size=8.5,
    )
    add_caption(doc, "Table 1. Evaluation method status. The actual-user methods are ready to field but not yet completed.")

    doc.add_heading("3. Data Snapshot", level=1)
    doc.add_picture(str(CHART_PATH), width=Inches(6.35))
    add_caption(doc, "Figure 1. Summary visualization from simulated proxy evaluation data, N=10. Not real survey statistics.")

    metric_rows = [
        ("Completed defensible methods", "3 / 7", "Expert/self evaluation, cognitive walkthrough, simulated user focus group."),
        ("Ready-to-field methods", "4 / 7", "Actual users, low-fi usability, survey, and high-fi usability protocols prepared."),
        ("Average ease rating", "3.4 / 5", "Users can likely complete the core flow."),
        ("Average confidence rating", "2.9 / 5", "Trust is moderate and needs stronger proof."),
        ("Average transparency rating", "2.5 / 5", "Evidence visibility and explanation are the biggest weaknesses."),
        ("Paper-only clarity requested", "10 / 10", "All proxy personas wanted clearer no-live-orders messaging."),
        ("Strategy evidence requested", "9 / 10", "Most proxy personas wanted DSR, PBO, benchmark, and report provenance in the UI."),
    ]
    add_simple_table(doc, ["Metric", "Value", "Meaning"], metric_rows, [2900, 1400, 5060], font_size=9)
    add_caption(doc, "Table 2. Headline data from the simulated proxy review and evaluation plan.")

    doc.add_heading("4. Evaluation Findings By Method", level=1)
    doc.add_heading("4.1 Expert / Self Evaluation", level=2)
    add_body_para(
        doc,
        "The expert review found that O(Alpha) is suitable as a controlled paper-trading prototype and research-to-paper workflow demo. It should not be presented as a live-trading product or as self-evidently validated without stronger in-product evidence surfaces.",
    )
    expert_rows = [
        ("Product scope", "4/5", "Research and paper trading scope is clearly safer than live trading.", "Repeat paper-only scope near risky-feeling actions."),
        ("Onboarding", "3.5/5", "Risk profile, strategy choice, backtest, and acceptance create a coherent setup path.", "Explain what the backtest does and does not prove."),
        ("Settings safety", "4.5/5", "Active-agent settings lock is strong and consistent with fail-closed safety.", "Keep backend as source of truth; add user-friendly reason copy."),
        ("Dashboard observability", "3/5", "Dashboard categories are right, but freshness and health state are not prominent enough.", "Add latest bar date, last snapshot, heartbeat, and stream state."),
        ("Strategy evidence", "2.5/5", "Validation evidence is mostly in docs/reports instead of user-facing cards.", "Add evidence drawer with DSR, PBO, benchmark, OOS trades, costs, and report path."),
    ]
    add_simple_table(doc, ["Feature", "Rating", "Analysis", "Suggested improvement"], expert_rows, [1800, 900, 3300, 3360], font_size=8.5)

    doc.add_heading("4.2 Cognitive Walkthrough / Heuristic Evaluation", level=2)
    walkthrough_rows = [
        ("Onboarding completion", "3.5/5", "Users can select a risk profile, choose a matching strategy, run a backtest, and accept it.", "Add beginner explanations for risk, Sharpe, Max DD, and historical simulation limits."),
        ("Agent launch / terminate", "4/5", "The main action is understandable and the lifecycle model is clear.", "Add paper-only badges beside Launch Agent and portfolio value."),
        ("Settings while active", "4.5/5", "Settings are locked while a portfolio agent is running, preventing confusing mid-run changes.", "Keep controls disabled and add concise reason text."),
        ("Evidence discovery", "2/5", "Users must leave the product to verify promotion artifacts and report lineage.", "Add report links and validation metadata directly to strategy cards."),
        ("Dashboard status", "3/5", "The dashboard shows portfolio categories but not enough stale/error/last-updated state.", "Add freshness, active run ID, heartbeat, and stream status."),
        ("Copy clarity", "3/5", "Some labels are polished but less direct than users need.", "Use plain labels such as Save settings, Saving, Settings locked, and Run backtest."),
    ]
    add_simple_table(doc, ["Task / heuristic", "Rating", "Analysis", "Suggested improvement"], walkthrough_rows, [2000, 900, 3200, 3260], font_size=8.5)

    doc.add_page_break()
    doc.add_heading("5. Simulated User / Persona Testing", level=1)
    add_callout(
        doc,
        "Evidence boundary",
        "The table below represents simulated proxy personas, not actual users. It is useful for planning product improvements and future research, but it should not be labelled as real user testing.",
        fill=PALE_RED,
    )
    add_body_para(
        doc,
        "Ratings below are overall averages across ease, confidence, and transparency signals.",
        after=4,
    )

    persona_rows = [
        ("Beginner paper trader", "Onboarding, dashboard, launch", "3.3/5", "Could complete onboarding and launch a paper agent, but was unsure about Sharpe, Max DD, h63, LGBM, and demo/live distinction.", "Add tooltips and visible Paper only / no brokerage orders labels."),
        ("Risk-aware retail investor", "Strategy proof, trades, settings", "2.7/5", "Liked safety boundaries but could not fully verify strategy quality from the UI.", "Add evidence-aware strategy cards and trade audit drawers."),
        ("Quant hobbyist", "Backtest evidence, catalog proof", "3.0/5", "Understood the workflow but wanted DSR, PBO, costs, benchmarks, and report links.", "Show Promote, DSR, PBO, benchmark, OOS trades, costs, and report path."),
        ("Skeptical product manager", "Onboarding narrative, proof", "3.0/5", "Saw a coherent product story but felt proof was mostly in docs rather than screens.", "Add a plain-language proof panel after onboarding and on the dashboard."),
        ("Reliability-focused engineer", "Lifecycle, health, failure state", "3.3/5", "Trusted the settings lock, but wanted more operational diagnostics.", "Add last refresh, active run ID, heartbeat, stream state, and failure reason."),
        ("Finance student", "Learning, backtest interpretation", "3.0/5", "Found onboarding approachable but might read Accept Backtest as this strategy is good.", "Change acceptance copy to acknowledge historical paper simulation limits."),
        ("Compliance-minded reviewer", "Claims, disclaimers, audit trail", "2.5/5", "Wanted claims to match artifact-backed evidence and paper-only limits.", "Add compliance-safe no-advice, no-live-order, and artifact lineage labels."),
        ("Active trader / power user", "Trade activity, exports, drilldowns", "2.8/5", "Wanted faster inspection of trades, alerts, and allocation changes.", "Add trade detail drawer with target weight, previous/current position, fill price, timestamp, and related alert."),
        ("Cautious nontechnical saver", "Safety, plain language", "2.7/5", "Liked guardrails but felt the product was technical and could over-trust the backtest.", "Add plain-language safety mode and clearer empty/demo states."),
        ("Privacy-conscious evaluator", "Seeded demo, reproducibility", "3.0/5", "Wanted to review without exposing personal financial data.", "Create deterministic demo script with seeded users and paper-only portfolios."),
    ]
    add_simple_table(
        doc,
        ["User / persona", "Feature tested", "How they felt", "Analysis", "Suggestions for improvement"],
        persona_rows,
        [1500, 1550, 950, 2800, 2560],
        font_size=7.8,
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("6. Theme Counts And Product Interpretation", level=1)
    theme_rows = [
        ("Settings lock felt reassuring", "10 / 10", "Keep the active-agent safety guardrail as a core product proof point."),
        ("Clearer paper-only messaging needed", "10 / 10", "Repeat no-live-orders language near every risky-feeling action."),
        ("Validation evidence should be visible in UI", "9 / 10", "Expose strategy proof rather than requiring users to inspect repo reports."),
        ("Quant terms or strategy names unclear", "8 / 10", "Add tooltips and beginner-mode explanations."),
        ("Freshness or health indicators needed", "7 / 10", "Make backend freshness, stream status, latest bar date, and heartbeat visible."),
        ("Trade rationale drilldowns needed", "6 / 10", "Explain why a simulated trade happened and how allocation changed."),
    ]
    add_simple_table(doc, ["Theme", "Proxy count", "Product interpretation"], theme_rows, [3300, 1200, 4860], font_size=9)

    doc.add_heading("7. Recommended Improvements", level=1)
    roadmap_rows = [
        ("P0", "Paper-only clarity", "Add persistent Paper only / no brokerage orders labels near Launch Agent, P&L, onboarding backtest acceptance, and settings.", "Reduces accidental over-trust and improves evaluator confidence."),
        ("P0", "Strategy evidence", "Add evidence-aware cards or drawers with promotion status, DSR, PBO, benchmark, OOS trades, cost stress, report path, and model artifact requirements.", "Brings the project's strongest validation story into the product UI."),
        ("P1", "Demo and empty states", "Replace realistic fallback values with explicit demo, unavailable, or empty states.", "Prevents missing backend state from looking like real performance."),
        ("P1", "Dashboard health", "Show latest bar date, last portfolio snapshot, last agent evaluation, active run ID, heartbeat, and stream connection state.", "Improves reliability trust and operational transparency."),
        ("P1", "Trade audit drawer", "Show triggering strategy, target weight change, previous/current position, simulated fill price, timestamp, and related alert.", "Answers the common why did this happen question."),
        ("P2", "Education layer", "Add tooltip explanations for Sharpe, Max DD, DSR, PBO, active sleeve, regime, LGBM, h63, leverage, stop-loss, and take-profit.", "Makes the product more approachable for beginner users."),
    ]
    add_simple_table(doc, ["Priority", "Area", "Recommendation", "Reason"], roadmap_rows, [850, 1650, 4050, 2810], font_size=8.5, header_fill=LIGHT_BLUE)

    doc.add_heading("8. Future Real-Participant Research Plan", level=1)
    future_rows = [
        ("Actual user focus group / interviews", "10 participants", "45-60 minute private sessions with seeded demo accounts; no real portfolio data.", "Not completed."),
        ("Low-fidelity usability testing", "5-8 participants", "Figma or wireframe tests for onboarding, evidence drawer, dashboard status, and settings lock.", "Not completed."),
        ("Survey of potential users", "50 respondents", "Private survey covering trust, safety expectations, strategy evidence, and feature priorities.", "Not fielded."),
        ("High-fidelity usability testing", "8-12 participants", "Working prototype with seeded account; measure task success, assists, trust, safety comprehension, and SUS.", "Not completed."),
    ]
    add_simple_table(doc, ["Method", "Target sample", "Privacy-safe setup", "Status"], future_rows, [2400, 1350, 4010, 1600], font_size=8.5)
    add_body_para(
        doc,
        "The best next evidence step is a private high-fidelity usability test with five real participants using seeded accounts. This would provide stronger evidence than additional simulated personas while avoiding public launch and protecting privacy.",
        before=6,
    )

    doc.add_page_break()
    doc.add_heading("9. Simulated Proxy Survey Results", level=1)
    add_callout(
        doc,
        "Evidence boundary",
        "The results below come from a simulated 50-person proxy survey run through subagents. They are useful as fast product-planning and investor-narrative signals, but they are not real survey responses and should not be represented as actual user research.",
        fill=PALE_RED,
    )
    add_body_para(
        doc,
        "The proxy survey was segmented into 20 beginner/intermediate potential users, 15 risk-aware or technical potential users, and 15 evaluator/compliance/finance-student personas. Ratings are weighted by segment size.",
    )
    simulated_survey_rows = [
        ("Concept clarity", "4.1 / 5", "The concept is easy to understand when framed as research-first paper trading."),
        ("Paper-only understanding", "4.4 / 5", "The paper-only positioning lands well and should be repeated visually in product."),
        ("Trust after safety guardrails", "4.1 / 5", "Active-agent locks and no-live-trading positioning are strong trust builders."),
        ("Desire for validation evidence", "4.7 / 5", "Report-backed strategy evidence is the strongest conversion driver."),
        ("Interest in beginner education", "4.2 / 5", "Guided explanations remain valuable, especially for beginner and student audiences."),
        ("Likelihood to try private demo", "3.9 / 5", "Private demo interest is positive but depends on clarity, simplicity, and proof."),
    ]
    add_simple_table(
        doc,
        ["Simulated survey metric", "Weighted result", "Interpretation"],
        simulated_survey_rows,
        [3000, 1600, 4760],
        font_size=8.8,
        header_fill=LIGHT_BLUE,
    )
    segment_rows = [
        ("Beginner/intermediate investing users", "20", "Concept clarity 4.1; paper-only 4.4; trust 4.2; validation evidence 4.6; education 4.5; demo likelihood 3.9."),
        ("Risk-aware / technical users", "15", "Concept clarity 4.1; paper-only 4.5; trust 4.3; validation evidence 4.8; education 3.7; demo likelihood 4.0."),
        ("Evaluator / compliance / finance-student users", "15", "Concept clarity 4.1; paper-only 4.4; trust 3.9; validation evidence 4.7; education 4.3; demo likelihood 3.8."),
    ]
    add_simple_table(
        doc,
        ["Proxy segment", "N", "Segment signal"],
        segment_rows,
        [3000, 700, 5660],
        font_size=8.5,
        header_fill=LIGHT_GRAY,
    )
    add_body_para(
        doc,
        "Strongest positive signal: the product is compelling when positioned as a paper-only research lab with visible safety controls and report-backed strategy evidence. Main funding-readiness gap: actual demo usage, real participant feedback, and credible proof that users return after the first session.",
        after=4,
    )
    add_bullets(
        doc,
        [
            "Simulated quote: \"The safety-lock framing makes it feel more like a research lab than a trading bot, which is the right direction.\"",
            "Simulated quote: \"Report-backed validation is the part I care about most. Show me the evidence and I'll spend time with it.\"",
            "Simulated quote: \"If it stays clearly paper-only and teaches the reasoning behind strategies, I can see this being useful for finance students without encouraging risky behavior.\"",
        ]
    )

    doc.add_heading("10. Potential User Survey Instrument", level=1)
    add_callout(
        doc,
        "Survey status",
        "A real potential-user survey has not been fielded yet. The instrument below is ready to use with 50 real potential users, but real-world results should only be reported after actual responses are collected and anonymized.",
        fill=PALE_YELLOW,
    )
    survey_rows = [
        ("Screener", "How would you describe your investing or trading experience?", "Beginner / intermediate / advanced / none", "Segments results by experience level."),
        ("Screener", "Have you used a paper-trading, investing, or backtesting tool before?", "Yes / no", "Separates familiar users from first-time users."),
        ("Concept clarity", "I understand that O(Alpha) is for paper trading only.", "1-5 agreement", "Checks whether the safety scope is understood."),
        ("Trust", "I would want to see validation evidence before launching an automated strategy.", "1-5 agreement", "Tests demand for report-backed proof."),
        ("Reverse-coded trust", "A five-year backtest alone is enough for me to trust a strategy.", "1-5 agreement", "Flags over-trust in simple backtests."),
        ("Safety concern", "I am concerned about accidentally placing real trades.", "1-5 agreement", "Measures need for stronger no-live-orders messaging."),
        ("Evidence", "I would find DSR, PBO, benchmark, and report links useful on strategy cards.", "1-5 agreement", "Validates the evidence-card roadmap."),
        ("Education", "I would use short explanations for terms like Sharpe, drawdown, PBO, and DSR.", "1-5 agreement", "Validates beginner education features."),
        ("Dashboard value", "Rank the most important dashboard information.", "Ranking", "Prioritizes P&L, positions, trades, alerts, strategy proof, and data freshness."),
        ("Open response", "What would make you trust this product more?", "Open text", "Captures unmet trust requirements."),
        ("Open response", "What would make you stop using this product?", "Open text", "Captures blockers and safety concerns."),
    ]
    add_simple_table(
        doc,
        ["Module", "Survey item", "Response type", "Why it matters"],
        survey_rows,
        [1400, 3800, 1800, 2360],
        font_size=8.2,
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("11. Honest Rating Positioning For Submission", level=1)
    add_body_para(
        doc,
        "The current ratings should remain as recorded. A stronger funding or submission story can still be made by separating current evidence from target outcomes after the recommended improvements. The target column below is a roadmap goal, not a claimed result.",
    )
    target_rows = [
        ("Ease", "3.4 / 5", "4.0+ / 5", "Add tooltips, clearer backtest acceptance copy, and beginner-mode explanations."),
        ("Confidence", "2.9 / 5", "4.0+ / 5", "Surface paper-only labels, proof panels, and report-backed strategy evidence."),
        ("Transparency", "2.5 / 5", "4.1+ / 5", "Add DSR, PBO, benchmark, OOS trades, costs, report path, and dashboard freshness indicators."),
        ("Simulated survey appeal", "4.1-4.7 / 5", "Validate with real users", "Field the prepared 50-person survey and compare real responses to the proxy benchmark."),
        ("Safety clarity", "10 / 10 proxy theme", "Maintain 10 / 10", "Keep active-agent lock, backend fail-closed guard, and explicit no-live-order messaging."),
    ]
    add_simple_table(
        doc,
        ["Metric", "Current evidence", "Target after improvements", "How to earn the target"],
        target_rows,
        [1500, 1800, 2000, 4060],
        font_size=8.5,
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("Appendix: Source Artifacts", level=1)
    add_bullets(
        doc,
        [
            "docs/evaluation/solution_suitability_evaluation_2026-07-05.md",
            "docs/evaluation/solution_suitability_visual_summary_2026-07-05.md",
            "docs/evaluation/solution_suitability_metrics_2026-07-05.csv",
            "docs/evaluation/solution_suitability_charts_2026-07-05.svg",
            "docs/evaluation/simulated_proxy_survey_results_2026-07-08.csv",
            "docs/evaluation/potential_user_survey_questionnaire_2026-07-08.csv",
            "README.md, docs/submission/OALPHA_README.md, onboarding, dashboard, settings, activity, and strategy catalog evidence files.",
        ]
    )

    doc.core_properties.title = "O(Alpha) Evaluation Summary and User Testing Readout"
    doc.core_properties.subject = "Solution suitability evaluation and simulated proxy user testing readout"
    doc.core_properties.author = "O(Alpha)"
    doc.core_properties.comments = "Private planning artifact. Simulated proxy data is clearly labelled."
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_doc()
    print(DOCX_PATH)
    print(CHART_PATH)

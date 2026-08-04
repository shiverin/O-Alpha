from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "evaluation"
DOCX_PATH = OUT_DIR / "OAlpha_User_Testing_and_Evaluation_Section_2026-07-11.docx"
FOCUS_CHART = OUT_DIR / "OAlpha_Simulated_Focus_Group_Ratings_2026-07-11.png"
SURVEY_CHART = OUT_DIR / "OAlpha_Simulated_Survey_Benchmark_2026-07-11.png"

INK = RGBColor(20, 20, 20)
MUTED = RGBColor(95, 95, 95)
WHITE = "FFFFFF"
LIGHT_GRAY = "E7E7E7"
MID_GRAY = "B8B8B8"
DARK_GRAY = "8A8A8A"
NOTE_GRAY = "F2F2F2"
WARNING_GRAY = "E0E0E0"


def set_run(run, *, size=10.5, bold=False, italic=False, color=INK, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=95, bottom=80, end=95):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def style_cell(cell, *, size=8.25, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.02
        for run in p.runs:
            set_run(run, size=size, bold=bold)


def add_table(doc, headers, rows, widths, *, size=8.25, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        shade(table.rows[0].cells[idx], header_fill)
        style_cell(table.rows[0].cells[idx], size=size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            style_cell(cells[idx], size=size)
    table_geometry(table, widths)
    return table


def add_body(doc, text, *, before=0, after=6, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if bold_lead and text.startswith(bold_lead):
        lead, rest = text[: len(bold_lead)], text[len(bold_lead) :]
        set_run(p.add_run(lead), bold=True)
        set_run(p.add_run(rest))
    else:
        set_run(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        if p.runs:
            p.runs[0].text = item
            set_run(p.runs[0], size=10)
        else:
            set_run(p.add_run(item), size=10)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(text), size=8.5, italic=True, color=MUTED)


def add_note(doc, title, text, *, fill=NOTE_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title), size=9.5, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.04
    set_run(p2.add_run(text), size=9.25)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def status_block(doc, status, evidence, sample):
    add_table(
        doc,
        ["Status", "Evidence base", "Sample / coverage"],
        [(status, evidence, sample)],
        [1700, 4860, 2800],
        size=8.5,
        header_fill=MID_GRAY,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in [
        ("Heading 1", 18, 16, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 11.5, 9, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("O(Alpha) | User Testing and Evaluation"), size=8, color=MUTED)
    set_run(footer.add_run("    "), size=8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_bar_chart(path, title, subtitle, rows, *, max_value=5.0):
    image = Image.new("RGB", (1600, 820), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 46), title, fill="#171717", font=font(38, True))
    draw.text((60, 98), subtitle, fill="#666666", font=font(20))
    x0, x1 = 660, 1450
    for idx, (label, value) in enumerate(rows):
        y = 182 + idx * 92
        draw.text((65, y), label, fill="#222222", font=font(23, True))
        draw.rounded_rectangle((x0, y, x1, y + 36), radius=8, fill="#ededed")
        fill_x = x0 + int((x1 - x0) * value / max_value)
        draw.rounded_rectangle((x0, y, fill_x, y + 36), radius=8, fill="#4b4b4b")
        draw.text((1470, y + 2), f"{value:.1f}", fill="#171717", font=font(22, True))
    draw.line((60, 748, 1540, 748), fill="#bdbdbd", width=2)
    draw.text((60, 770), "Scale: 1 = low, 5 = high. Values are simulated proxy data, not real participant statistics.", fill="#666666", font=font(17))
    image.save(path)


def build_charts():
    draw_bar_chart(
        FOCUS_CHART,
        "Simulated focus group: average ratings",
        "Ten proxy personas reviewed the current product flow and evidence story.",
        [("Ease of core workflow", 3.4), ("Confidence in the product", 2.9), ("Transparency of evidence", 2.5)],
    )
    draw_bar_chart(
        SURVEY_CHART,
        "Simulated survey: product-planning benchmark",
        "Weighted proxy benchmark across 50 simulated responses in three target segments.",
        [
            ("Concept clarity", 4.1),
            ("Paper-only understanding", 4.4),
            ("Trust after safety guardrails", 4.1),
            ("Demand for validation evidence", 4.7),
            ("Interest in beginner education", 4.2),
            ("Likelihood to try a private demo", 3.9),
        ],
    )


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_charts()
    doc = Document()
    setup_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(110)
    title.paragraph_format.space_after = Pt(10)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(title.add_run("User Testing and Evaluation"), size=24, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    set_run(subtitle.add_run("O(Alpha) solution suitability assessment"), size=12.5, color=MUTED)
    add_note(
        doc,
        "Evidence integrity",
        "Completed evidence includes expert/self evaluation, a cognitive walkthrough, and simulated proxy reviews. Actual interviews and real-participant low-fidelity, survey, and high-fidelity studies have not yet been conducted. Protocols are documented so those studies can be run privately with seeded accounts and no personal financial data.",
        fill=WARNING_GRAY,
    )
    add_body(
        doc,
        "This section evaluates whether O(Alpha) is understandable, safe, and suitable as a paper-trading and quantitative research prototype. Each evaluation method is presented separately using a consistent structure so that completed evidence, simulated evidence, and future research are easy to distinguish.",
        before=14,
    )
    doc.add_heading("Evaluation Method Overview", level=2)
    overview_rows = [
        ("Expert / self evaluation", "Completed", "Yes"),
        ("Cognitive walkthrough / heuristic evaluation", "Completed", "Yes"),
        ("Simulated user focus group", "Completed - simulated", "Yes, labelled simulated"),
        ("Actual user focus group / interview", "Protocol prepared", "No"),
        ("Low-fidelity usability testing", "Protocol prepared", "No"),
        ("Survey of potential users", "Proxy benchmark only", "No real survey"),
        ("High-fidelity usability testing", "Protocol prepared", "No"),
    ]
    add_table(doc, ["Evaluation method", "Current status", "Claim as completed?"], overview_rows, [5000, 2440, 1920], size=8.8)

    doc.add_page_break()
    doc.add_heading("1. Expert / Self Evaluation", level=1)
    status_block(doc, "Completed", "Current repository, product flows, safety controls, and technical documentation", "Product-level review")
    doc.add_heading("Objective and approach", level=3)
    add_body(doc, "The evaluation assessed whether O(Alpha) is suitable as a controlled paper-trading and quantitative research product for beginner users, risk-aware users, evaluators, and technical reviewers. The review covered onboarding, strategy selection, backtest acceptance, dashboard observability, agent lifecycle, settings safety, and strategy evidence.")
    doc.add_heading("Evaluation results", level=3)
    expert_rows = [
        ("Product scope", "4.0 / 5", "Research and paper-trading boundaries are appropriate.", "Repeat paper-only scope beside high-risk actions."),
        ("Onboarding", "3.5 / 5", "Risk selection, strategy choice, backtest, and acceptance form a coherent setup path.", "Explain what the backtest does and does not prove."),
        ("Settings safety", "4.5 / 5", "The active-agent settings lock is strong and consistent with fail-closed safety.", "Retain the backend guard and explain the lock in plain language."),
        ("Dashboard observability", "3.0 / 5", "Core portfolio categories are present, but freshness and health are not prominent.", "Show latest bar, snapshot, heartbeat, run ID, and stream state."),
        ("Strategy evidence", "2.5 / 5", "Validation evidence is stronger in reports than in the interface.", "Add DSR, PBO, benchmark, out-of-sample trades, costs, and report provenance."),
    ]
    add_table(doc, ["Feature", "Rating", "Analysis", "Suggested improvement"], expert_rows, [1700, 1000, 3250, 3410], size=8.2)
    add_caption(doc, "Table 1. Expert/self evaluation results. Ratings use a five-point suitability scale.")
    doc.add_heading("Conclusion", level=3)
    add_body(doc, "O(Alpha) is suitable as a controlled paper-trading prototype and research-to-paper workflow demonstration. Its strongest product proof is the active-agent safety boundary. Its principal suitability gap is trust calibration: the interface should surface stronger evidence, clearer limits, and more operational status before the product is presented beyond a prototype context.")

    doc.add_page_break()
    doc.add_heading("2. Cognitive Walkthrough / Heuristic Evaluation", level=1)
    status_block(doc, "Completed", "Task-by-task inspection of onboarding, dashboard, agent controls, settings, and evidence discovery", "10-step core journey")
    doc.add_heading("Walkthrough tasks", level=3)
    add_bullets(doc, [
        "Create or sign in to an account and complete onboarding.",
        "Choose a risk profile and a compatible catalog strategy.",
        "Run and accept the onboarding backtest.",
        "Identify portfolio status and launch the paper portfolio agent.",
        "Inspect allocation, positions, executions, alerts, and activity.",
        "Attempt to edit settings while the agent is active.",
        "Stop the agent and save a settings change.",
        "Locate evidence supporting the selected strategy.",
    ])
    doc.add_heading("Heuristic findings", level=3)
    heuristic_rows = [
        ("System status", "Medium", "Active state is visible; freshness, last evaluation, and stream state are less visible.", "Add last updated, heartbeat, latest bar date, active run ID, and connection status."),
        ("Real-world language", "Medium", "Finance terms are accurate but difficult for beginners.", "Add plain-language definitions for Sharpe, drawdown, PBO, DSR, regime, and cadence."),
        ("Control and freedom", "Low", "Launch, terminate, and settings-lock behavior form a clear safety model.", "Keep the lock and state why mid-run settings cannot change."),
        ("Error prevention", "Low", "Frontend and backend settings guardrails prevent unsafe changes.", "Retain backend enforcement as the source of truth."),
        ("Recognition", "High", "Users must leave the product to inspect strategy evidence and provenance.", "Add evidence drawers and report links to strategy cards."),
        ("Trust", "High", "Paper-only scope is not repeated enough near actions that feel financially consequential.", "Place Paper only / no brokerage orders labels beside launch, P&L, backtest acceptance, and settings."),
        ("Fallback states", "High", "Realistic fallback values may appear to be genuine portfolio performance.", "Use explicit demo, unavailable, or empty states."),
    ]
    add_table(doc, ["Heuristic", "Severity", "Finding", "Recommendation"], heuristic_rows, [1500, 900, 3350, 3610], size=8.15)
    add_caption(doc, "Table 2. Cognitive walkthrough and heuristic evaluation findings.")
    doc.add_heading("Conclusion", level=3)
    add_body(doc, "The core path is coherent and safety-aware. Users can complete the intended actions, but confidence depends on understanding what is simulated, seeing proof for strategy claims, and knowing whether system data is current.")

    doc.add_page_break()
    doc.add_heading("3. Simulated User Focus Group", level=1)
    status_block(doc, "Completed - simulated", "Ten proxy personas grounded in the current product and documentation", "N = 10 simulated personas")
    add_note(doc, "Evidence boundary", "This is a simulated proxy focus group, not an actual focus group and not real participant research. The results support product planning and research design only.", fill=WARNING_GRAY)
    doc.add_heading("Method", level=3)
    add_body(doc, "Ten distinct proxy personas assessed the onboarding flow, paper-agent lifecycle, settings safety, dashboard clarity, strategy evidence, and product trust. Each persona provided ease, confidence, and transparency reactions on a five-point scale together with a primary concern and suggested improvement.")
    doc.add_picture(str(FOCUS_CHART), width=Inches(6.55))
    add_caption(doc, "Figure 1. Average ratings from the simulated 10-person proxy focus group.")
    doc.add_heading("Persona-level results", level=3)
    persona_rows = [
        ("Beginner paper trader", "Onboarding and launch", "3.3 / 5", "Core path was usable; quant terms and paper/live distinction were unclear.", "Add tooltips and persistent paper-only labels."),
        ("Risk-aware retail investor", "Strategy proof and settings", "2.7 / 5", "Safety boundaries were reassuring, but strategy quality was difficult to verify.", "Add evidence-aware strategy cards and trade audit detail."),
        ("Quant hobbyist", "Backtest evidence", "3.0 / 5", "Wanted deeper validation metadata and report lineage.", "Show DSR, PBO, benchmark, costs, out-of-sample trades, and report path."),
        ("Skeptical product manager", "Product story and proof", "3.0 / 5", "The flow was coherent, but proof remained concentrated in documentation.", "Add a plain-language proof panel in the product."),
        ("Reliability engineer", "Lifecycle and health", "3.3 / 5", "Trusted the settings lock; wanted clearer operational diagnostics.", "Show heartbeat, stream state, run ID, last refresh, and failure reason."),
        ("Finance student", "Learning and backtests", "3.0 / 5", "Onboarding felt approachable, but accepting a backtest could imply endorsement.", "Use explicit historical-simulation acknowledgement copy."),
        ("Compliance reviewer", "Claims and auditability", "2.5 / 5", "Wanted product claims to match artifact-backed evidence and paper-only limits.", "Add no-advice, no-live-order, and provenance labels."),
        ("Active trader", "Trades and drilldowns", "2.8 / 5", "Wanted faster inspection of simulated fills and allocation changes.", "Add a trade drawer with rationale, weight change, fill price, and alert."),
        ("Cautious nontechnical saver", "Safety and language", "2.7 / 5", "Liked guardrails but could over-trust realistic values and backtests.", "Add beginner explanations and explicit demo states."),
        ("Privacy-conscious evaluator", "Seeded demo", "3.0 / 5", "Wanted a reproducible review without personal financial data.", "Provide deterministic seeded accounts and sample portfolios."),
    ]
    add_table(doc, ["User / persona", "Feature", "How they felt", "Analysis", "Suggestion"], persona_rows, [1500, 1400, 1050, 2850, 2560], size=7.6)
    add_caption(doc, "Table 3. Simulated persona feedback. The rating combines ease, confidence, and transparency signals.")
    doc.add_heading("Cross-persona themes", level=3)
    theme_rows = [
        ("Settings lock felt reassuring", "10 / 10", "Keep the guardrail as a visible product proof point."),
        ("Clearer paper-only messaging needed", "10 / 10", "Repeat no-live-order language near consequential actions."),
        ("Validation evidence should be visible", "9 / 10", "Expose strategy proof in the interface."),
        ("Quant terms were unclear", "8 / 10", "Add tooltips and beginner explanations."),
        ("Freshness or health indicators needed", "7 / 10", "Show data age, heartbeat, stream state, and latest evaluation."),
        ("Trade rationale drilldowns needed", "6 / 10", "Explain why each simulated action occurred."),
    ]
    add_table(doc, ["Theme", "Proxy count", "Interpretation"], theme_rows, [3400, 1300, 4660], size=8.4)

    doc.add_heading("4. Actual User Focus Group / Interview", level=1)
    status_block(doc, "Not yet conducted", "Private interview protocol prepared", "Target N = 10 real participants")
    add_note(doc, "Reporting rule", "Do not claim this method as completed until real participants have taken part and anonymized notes have been retained.")
    doc.add_heading("Purpose and participant mix", level=3)
    add_body(doc, "The proposed study tests whether potential users understand the paper-only scope, can complete the core workflow, and can explain what evidence they require before trusting an automated strategy. A balanced sample should include beginner investors, risk-aware retail investors, technically inclined users, a product evaluator, a compliance or risk reviewer, and a finance learner.")
    doc.add_heading("Private session procedure", level=3)
    add_bullets(doc, [
        "Run 45-60 minute moderated sessions using seeded demo accounts.",
        "Collect no brokerage credentials, real portfolio data, or personal financial-position details.",
        "Ask participants to choose a risk profile, run and interpret a backtest, launch the paper agent, inspect activity, test the settings lock, stop the agent, and locate strategy evidence.",
        "Record anonymized task outcomes, confusion points, trust and safety ratings, and improvement suggestions.",
    ])
    doc.add_heading("Interview record template", level=3)
    interview_rows = [
        ("P01", "Segment", "Completed / assisted / failed", "Observation", "1-5", "1-5", "Recommendation"),
        ("P02", "Segment", "Completed / assisted / failed", "Observation", "1-5", "1-5", "Recommendation"),
        ("P03-P10", "Repeat per participant", "-", "-", "-", "-", "-"),
    ]
    add_table(doc, ["Participant", "Segment", "Task success", "Main confusion", "Trust", "Safety", "Suggestion"], interview_rows, [1000, 1200, 1800, 1900, 800, 800, 1860], size=7.9)
    add_caption(doc, "Table 4. Anonymized interview record structure for future real-participant sessions.")
    doc.add_heading("Success criteria", level=3)
    add_bullets(doc, [
        "At least 80% correctly identify that the system is paper-only.",
        "At least 80% complete onboarding without moderator intervention.",
        "All participants observe that settings cannot be changed during an active run.",
        "At least 70% locate strategy evidence or clearly state where they expected it.",
    ])

    doc.add_heading("5. Usability Testing With Potential Users on Low-Fidelity Artefacts", level=1)
    status_block(doc, "Not yet conducted", "Figma or static-wireframe protocol prepared", "Target N = 5-8 real participants")
    add_note(doc, "Reporting rule", "A low-fidelity protocol is research preparation, not evidence that potential users have tested the artefacts.")
    doc.add_heading("Artefacts and research questions", level=3)
    add_body(doc, "The study should use low-fidelity screens for onboarding, strategy evidence, dashboard health, and the locked-settings state. It is intended to test information architecture and comprehension before additional interface polish is implemented.")
    lowfi_rows = [
        ("Onboarding flow", "Can users identify where to begin and select a risk profile?", "Completion without help", "80%+"),
        ("Strategy card", "Can users identify a compatible strategy and find validation evidence?", "Correct selection and evidence discovery", "70%+"),
        ("Dashboard status", "Can users identify whether data is current and whether an agent is active?", "Correct state explanation", "80%+"),
        ("Settings lock", "Can users predict why controls are disabled during a run?", "Correct safety explanation", "90%+"),
        ("Paper-only scope", "Can users state whether any real money is involved?", "Correct scope comprehension", "90%+"),
    ]
    add_table(doc, ["Artefact", "Research question", "Metric", "Target"], lowfi_rows, [1900, 4080, 2360, 1020], size=8.35)
    add_caption(doc, "Table 5. Proposed low-fidelity usability study matrix.")
    doc.add_heading("Session tasks", level=3)
    add_bullets(doc, [
        "Point to where you would begin and choose a risk profile.",
        "Identify the safest compatible strategy and find proof supporting it.",
        "Explain what happens when settings are edited during an active agent run.",
        "Identify whether the prototype can place live brokerage orders.",
        "Describe one screen element that would increase trust.",
    ])

    doc.add_heading("6. Survey of Potential Users", level=1)
    status_block(doc, "Not fielded with real users", "Survey instrument plus simulated proxy benchmark", "Proxy N = 50; real N = 0")
    add_note(doc, "Evidence boundary", "The chart and ratings below are simulated proxy responses. They are a planning benchmark, not the result of a survey of potential users and should not be reported as real respondent data.", fill=WARNING_GRAY)
    doc.add_heading("Simulated benchmark", level=3)
    doc.add_picture(str(SURVEY_CHART), width=Inches(6.55))
    add_caption(doc, "Figure 2. Weighted simulated survey benchmark across three proxy segments.")
    survey_rows = [
        ("Beginner / intermediate investors", "20", "4.1", "4.4", "4.2", "4.6", "4.5", "3.9"),
        ("Risk-aware technical users", "15", "4.1", "4.5", "4.3", "4.8", "3.7", "4.0"),
        ("Evaluator / compliance / finance learners", "15", "4.1", "4.4", "3.9", "4.7", "4.3", "3.8"),
        ("Weighted aggregate", "50", "4.1", "4.4", "4.1", "4.7", "4.2", "3.9"),
    ]
    add_table(doc, ["Proxy segment", "N", "Clarity", "Paper-only", "Trust", "Evidence", "Education", "Private demo"], survey_rows, [2600, 500, 900, 1000, 800, 900, 1000, 1660], size=7.8)
    add_caption(doc, "Table 6. Simulated proxy survey results on a five-point scale.")
    doc.add_heading("Interpretation", level=3)
    add_body(doc, "The strongest planning signal is demand for visible validation evidence (4.7/5), followed by understanding of the paper-only positioning (4.4/5). Private-demo intent is positive but conditional (3.9/5), indicating that clarity, evidence, and simplicity are likely prerequisites for conversion. These hypotheses require validation with real respondents.")
    doc.add_heading("Ready-to-field survey instrument", level=3)
    instrument_rows = [
        ("Experience", "How would you describe your investing or trading experience?", "Category"),
        ("Product familiarity", "Have you used a paper-trading or backtesting tool before?", "Yes / No"),
        ("Scope clarity", "I understand that O(Alpha) is for paper trading only.", "1-5 agreement"),
        ("Trust", "I would want validation evidence before launching an automated strategy.", "1-5 agreement"),
        ("Over-trust check", "A five-year backtest alone is enough for me to trust a strategy.", "1-5 reverse-coded"),
        ("Safety", "I am concerned about accidentally placing real trades.", "1-5 agreement"),
        ("Evidence", "I would find DSR, PBO, benchmark, and report links useful.", "1-5 agreement"),
        ("Education", "I would use short explanations for unfamiliar finance terms.", "1-5 agreement"),
        ("Feature priority", "Rank P&L, positions, trades, alerts, strategy proof, and freshness.", "Ranking"),
        ("Open response", "What would make you trust this product more?", "Text"),
    ]
    add_table(doc, ["Module", "Survey item", "Response"], instrument_rows, [1800, 5940, 1620], size=8.15)
    add_caption(doc, "Table 7. Proposed questionnaire for a future private survey of potential users.")

    doc.add_heading("7. Usability Testing With Potential Users on High-Fidelity Artefacts", level=1)
    status_block(doc, "Not yet conducted with real users", "Working-prototype protocol prepared", "Target N = 8-12 real participants")
    add_note(doc, "Reporting rule", "Automated and simulated tests can validate behavior, but they do not replace observation of real potential users interacting with the working prototype.")
    doc.add_heading("Test environment", level=3)
    add_bullets(doc, [
        "Use a local or private deployment with deterministic seeded accounts and paper portfolios.",
        "Disable brokerage integration and collect no personal financial data.",
        "Log only task events, errors, timings, assists, and anonymized ratings.",
        "Reset the test account before each session to keep results comparable.",
    ])
    doc.add_heading("Task and measurement plan", level=3)
    highfi_rows = [
        ("Complete onboarding", "Success, time, assists", "80%+ unassisted"),
        ("Run and interpret a backtest", "Success, interpretation accuracy, SEQ", "80%+ identify historical limits"),
        ("Launch the paper agent", "Success, paper-only comprehension", "90%+ identify no real money at risk"),
        ("Inspect portfolio state", "Find value, allocation, alerts, and trades", "80%+ locate all four"),
        ("Attempt settings change while active", "Observed prevention and explanation", "100% blocked; 90% explain why"),
        ("Stop agent and save settings", "Success, errors, backtracks", "90%+ complete"),
        ("Find strategy evidence", "Discovery or correctly stated expectation", "70%+"),
        ("Post-session assessment", "SUS, trust before/after, safety comprehension", "SUS 75+; no decline in safety comprehension"),
    ]
    add_table(doc, ["Task", "Measures", "Pass criterion"], highfi_rows, [3200, 3440, 2720], size=8.25)
    add_caption(doc, "Table 8. High-fidelity working-prototype usability test plan.")
    doc.add_heading("Analysis and reporting", level=3)
    add_body(doc, "Report task success, time-on-task distributions, moderator assists, errors, backtracks, Single Ease Question ratings, System Usability Scale results, and pre/post trust. Segment results by beginner and experienced participants. Use anonymized quotes to explain patterns, and do not claim statistical significance from a small usability sample.")

    doc.add_heading("Consolidated Findings and Next Steps", level=1)
    add_body(doc, "The completed evaluations agree on one central finding: O(Alpha)'s workflow is understandable and its settings lock is reassuring, but users need stronger evidence visibility, clearer paper-only language, and clearer operational status before confidence can match usability.")
    priority_rows = [
        ("P0", "Paper-only clarity", "Add persistent Paper only / no brokerage orders labels beside launch, P&L, backtest acceptance, and settings."),
        ("P0", "Strategy evidence", "Expose promotion status, DSR, PBO, benchmark, out-of-sample trades, costs, report path, and model-artifact requirements."),
        ("P1", "System freshness", "Show latest bar, portfolio snapshot, agent evaluation, run ID, heartbeat, and stream state."),
        ("P1", "Demo and empty states", "Replace realistic fallback values with explicit demo, unavailable, or empty states."),
        ("P1", "Trade auditability", "Add rationale, target-weight change, previous/current position, simulated fill price, timestamp, and related alert."),
        ("P2", "Education", "Explain Sharpe, drawdown, DSR, PBO, regime, LGBM, h63, leverage, stop-loss, and take-profit in plain language."),
    ]
    add_table(doc, ["Priority", "Area", "Action"], priority_rows, [900, 1900, 6560], size=8.6, header_fill=MID_GRAY)
    add_caption(doc, "Table 9. Prioritized improvements derived from completed and simulated evaluations.")
    doc.add_heading("Recommended next evidence", level=3)
    add_body(doc, "The highest-value next step is a private high-fidelity usability test with five real participants using seeded accounts. It provides stronger product evidence than additional simulated personas while preserving privacy and avoiding public launch. The 50-person survey should follow once the prototype's messaging and evidence surfaces have been improved.")
    doc.core_properties.title = "O(Alpha) User Testing and Evaluation"
    doc.core_properties.subject = "Solution suitability evaluation and future user research protocols"
    doc.core_properties.author = "O(Alpha)"
    doc.core_properties.comments = "Simulated evidence is clearly labelled; actual participant methods are not claimed as completed."
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
